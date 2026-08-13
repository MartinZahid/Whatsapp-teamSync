# -*- coding: utf-8 -*-
import os
import tempfile
from datetime import date
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.dates import MonthLocator, DateFormatter, WeekdayLocator, FR
from matplotlib.patches import Patch
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOC_OUTPUT = "Reporte-Preliminar-WhatsApp-Team-Sync.docx"
RUTA_LOGO = None  # Cambiar por la ruta del logo de ITH, ej. "logo_ith.png"

FECHA_INICIO_PROYECTO = date(2026, 8, 7)     # viernes 7 de agosto de 2026
FECHA_FIN_PROYECTO = date(2026, 12, 10)      # jueves 10 de diciembre de 2026

NOMBRE_ESTUDIANTE = "[Nombre completo del estudiante]"
NUMERO_CONTROL = "[Número de control]"
CORREO_ESTUDIANTE = "[correo@estudiante.com]"
TELEFONO_ESTUDIANTE = "[Teléfono]"
CARRERA = "Ingeniería en Sistemas Computacionales"
ASESOR_EXTERNO = "[Nombre del asesor externo]"
FECHA = "[DD/MM/AAAA]"

TITULO = ("WhatsApp Team Sync: sistema de presencia en tiempo real y portal "
          "de monitoreo para equipos de atención al cliente en WhatsApp Web")

EMPRESA = {
    "nombre": "[Nombre de la empresa / organismo / dependencia]",
    "direccion": "[Dirección]",
    "telefono": "[Teléfono]",
    "correo": "[correo@empresa.com]",
    "giro": "[Giro de la empresa]",
    "ciudad": "[Ciudad, Estado]",
}

LUGAR_PROYECTO = "[Departamento o ubicación geográfica donde se realizará el proyecto]"

ASESOR = {
    "departamento": "[Departamento adscrito]",
    "puesto": "[Puesto]",
    "correo": "[correo@asesor.com]",
    "telefono": "[Teléfono]",
}

FECHA_INICIO = "07/08/2026"
FECHA_FIN = "10/12/2026"

# (actividad, fecha_inicio, fecha_fin, fase) — cronograma de 4 meses (18 semanas,
# cada actividad de una semana, consecutivas sin dejar días vacíos)
ACTIVIDADES_CRONOGRAMA = [
    ("1. Revisión del estado del arte y análisis de soluciones similares (WhatsApp Business API, CRM de mensajería)",
     date(2026, 8, 7), date(2026, 8, 13), "Análisis"),
    ("2. Levantamiento de requerimientos con el equipo de atención al cliente",
     date(2026, 8, 14), date(2026, 8, 20), "Análisis"),
    ("3. Diseño de la arquitectura de la solución (extensión, service worker, servidor, portal web) y del protocolo",
     date(2026, 8, 21), date(2026, 8, 27), "Diseño y configuración"),
    ("4. Configuración del entorno de desarrollo (Vite, TypeScript, Node.js, Git)",
     date(2026, 8, 28), date(2026, 9, 3), "Diseño y configuración"),
    ("5. Desarrollo del servidor WebSocket (presencia, rooms, heartbeat, rate limiting)",
     date(2026, 9, 4), date(2026, 9, 10), "Desarrollo"),
    ("6. Desarrollo de la extensión de Chrome: popup de configuración y página de opciones",
     date(2026, 9, 11), date(2026, 9, 17), "Desarrollo"),
    ("7. Desarrollo del content script: panel flotante en Shadow DOM y detector de contacto activo",
     date(2026, 9, 18), date(2026, 9, 24), "Desarrollo"),
    ("8. Implementación de la detección automática de chats y sincronización entre pestañas",
     date(2026, 9, 25), date(2026, 10, 1), "Desarrollo"),
    ("9. Desarrollo de la base de datos SQLite (sesiones de atención y bitácora de eventos)",
     date(2026, 10, 2), date(2026, 10, 8), "Desarrollo"),
    ("10. Desarrollo del portal web de monitoreo (En vivo, Historial, Reportes con Chart.js y Gestión)",
     date(2026, 10, 9), date(2026, 10, 15), "Desarrollo"),
    ("11. Integración del protocolo y pruebas de conexión, reconexión, estados y solicitudes de ayuda",
     date(2026, 10, 16), date(2026, 10, 22), "Integración y pruebas"),
    ("12. Pruebas unitarias y corrección de errores detectados",
     date(2026, 10, 23), date(2026, 10, 29), "Integración y pruebas"),
    ("13. Pruebas funcionales con un equipo real de atención al cliente",
     date(2026, 10, 30), date(2026, 11, 5), "Integración y pruebas"),
    ("14. Pruebas de rendimiento y carga del servidor",
     date(2026, 11, 6), date(2026, 11, 12), "Integración y pruebas"),
    ("15. Capacitación a usuarios del sistema",
     date(2026, 11, 13), date(2026, 11, 19), "Documentación y cierre"),
    ("16. Elaboración de manuales de instalación y uso",
     date(2026, 11, 20), date(2026, 11, 26), "Documentación y cierre"),
    ("17. Despliegue del sistema y puesta en producción",
     date(2026, 11, 27), date(2026, 12, 3), "Documentación y cierre"),
    ("18. Documentación técnica, empacado del sistema y elaboración del reporte final",
     date(2026, 12, 4), date(2026, 12, 10), "Documentación y cierre"),
]

COLORES_FASE = {
    "Análisis": "#3b82f6",
    "Diseño y configuración": "#8b5cf6",
    "Desarrollo": "#22c55e",
    "Integración y pruebas": "#f59e0b",
    "Documentación y cierre": "#ef4444",
}


def formato_fecha(f):
    return f.strftime("%d/%m/%Y")


def build_gantt(actividades, colores, salida):
    acts = actividades
    fig, ax = plt.subplots(figsize=(9.5, 5.5), dpi=200)
    for i, (_, ini, fin, fase) in enumerate(acts):
        ax.barh(i, (fin - ini).days + 1, left=ini, height=0.55,
                color=colores[fase], edgecolor="black", linewidth=0.4,
                align="center")
    ax.set_yticks(range(len(acts)))
    ax.set_yticklabels([a[0] for a in acts], fontsize=7)
    ax.xaxis.set_major_locator(MonthLocator())
    ax.xaxis.set_major_formatter(DateFormatter("%b\n%Y"))
    ax.xaxis.set_minor_locator(WeekdayLocator(byweekday=FR))
    ax.grid(axis="x", which="major", linestyle="--", alpha=0.5)
    ax.invert_yaxis()
    ax.set_xlim(FECHA_INICIO_PROYECTO, FECHA_FIN_PROYECTO)
    ax.set_xlabel("")
    handles = [Patch(color=c, label=f) for f, c in colores.items()]
    ax.legend(handles=handles, loc="upper right", fontsize=8, frameon=True)
    ax.set_title("Cronograma preliminar de actividades", fontsize=11, pad=12)
    fig.tight_layout()
    fig.savefig(salida, bbox_inches="tight")
    plt.close(fig)

RECURSOS = [
    ("Computadora portátil o de escritorio (Node.js 18+)", "Hardware", 1, "[Costo]"),
    ("Google Chrome (versión reciente)", "Software", 1, "0"),
    ("Cuenta de WhatsApp Web", "Software", 1, "0"),
    ("Conexión a internet", "Servicio", 1, "[Costo]"),
    ("Servidor para el servicio WebSocket y portal web (local o en la nube)", "Hardware/Software", 1, "[Costo]"),
    ("Entorno de desarrollo (Visual Studio Code, Git, Node.js, TypeScript)", "Software", 1, "0"),
    ("Base de datos SQLite (sql.js) y librería de gráficas Chart.js", "Software", 1, "0"),
]

recursos_totales = "[Costo total estimado]"


def set_font(run, size=11, bold=False, italic=False, color=None, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def add_para(doc, text="", size=11, bold=False, italic=False, align=None, color=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), size=size, bold=True, color=(0x1F, 0x3B, 0x73))
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pPr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6",
        qn("w:space"): "1", qn("w:color"): "1F3B73"})
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_numbered(doc, items, size=11):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(f"{i}. "), size=size, bold=True)
        set_font(p.add_run(item), size=size)


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(item), size=size)


def make_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(h), size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = hdr[j]._element.get_or_add_tcPr().makeelement(qn("w:shd"), {})
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F3B73")
        hdr[j]._element.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(str(val)), size=10)
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================ a) PORTADA ============================
add_para(doc, "", space_after=12)
add_para(doc, "[LOGO DEL INSTITUTO TECNOLÓGICO DE HERMOSILLO]",
         size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "INSTITUTO TECNOLÓGICO DE HERMOSILLO",
         size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Tecnológico Nacional de México",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para(doc, "REPORTE PRELIMINAR", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "Residencia Profesional", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

add_para(doc, "Título del proyecto", size=11, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
set_font(p.add_run(TITULO), size=13, bold=True)

datos_portada = [
    ("Nombre del estudiante", NOMBRE_ESTUDIANTE),
    ("Carrera", CARRERA),
    ("Número de control", NUMERO_CONTROL),
    ("Correo electrónico", CORREO_ESTUDIANTE),
    ("Teléfono", TELEFONO_ESTUDIANTE),
    ("Asesor externo", ASESOR_EXTERNO),
]
for label, valor in datos_portada:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(f"{label}: "), size=12, bold=True)
    set_font(p.add_run(valor), size=12)

add_para(doc, "", space_after=12)
add_para(doc, "Lugar y fecha: Hermosillo, Sonora, a " + FECHA,
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Firma del asesor externo"), size=11, italic=True)
doc.add_page_break()

# ======================= b) OBJETIVO GENERAL ========================
add_heading(doc, "b) Objetivo general del proyecto", size=14)
add_para(doc, ("Desarrollar un sistema integral de presencia en tiempo real para "
               "equipos de atención al cliente que trabajan sobre WhatsApp Web, "
               "compuesto por tres componentes: (1) una extensión para Google Chrome "
               "que detecta qué chat atiende cada agente y muestra su estado en un "
               "panel flotante, (2) un servidor de comunicaciones en Node.js que "
               "centraliza la presencia mediante WebSocket y persiste la información "
               "en una base de datos, y (3) un portal web de monitoreo (dashboard) que "
               "permita ver en tiempo real el estado del equipo y generar métricas de "
               "desempeño, con el fin de mejorar la coordinación, reducir los tiempos "
               "de respuesta y evitar la duplicidad de atención."))

# ============== c) DELIMITACIÓN Y RESUMEN DEL PROYECTO ==============
add_heading(doc, "c) Delimitación y resumen del proyecto", size=14)
add_para(doc, "Delimitación del proyecto:", bold=True)
add_bullets(doc, [
    "El sistema se ejecuta únicamente como extensión de Google Chrome (Manifest V3) sobre la interfaz de WhatsApp Web (web.whatsapp.com).",
    "La extensión incorpora un panel flotante con Shadow DOM que no interfiere con el funcionamiento normal de WhatsApp Web.",
    "Se contemplan cuatro estados de presencia por agente: Activo (atendiendo), Disponible, Pausado y Desconectado.",
    "El servidor de comunicaciones se desarrolla en Node.js con WebSocket y mantiene el registro de agentes y su presencia.",
    "El portal web (dashboard) es servido por el propio servidor e incluye las secciones: En vivo, Historial, Reportes y Gestión.",
    "La información se persiste en una base de datos SQLite (sesiones de chat y bitácora de eventos) y puede exportarse en formato JSON.",
    "Quedan fuera del alcance: la integración con WhatsApp Business API, el envío automatizado o masivo de mensajes, y versiones móviles.",
])
add_para(doc, "Resumen del proyecto:", bold=True)
add_para(doc, ("WhatsApp Team Sync es un sistema de presencia en tiempo real dirigido a "
               "equipos de servicio al cliente que trabajan sobre WhatsApp Web. Se "
               "compone de una extensión de navegador que detecta automáticamente el "
               "chat que abre cada agente e informa su estado (activo, disponible o "
               "pausado) a un servidor central mediante WebSocket; un servidor Node.js "
               "que administra la presencia, mantiene la sesión de cada agente, "
               "reconecta perfiles ante fallas de red y persiste en SQLite las sesiones "
               "de atención y la bitácora de eventos; y un portal web de monitoreo que "
               "muestra en tiempo real quién atiende cada conversación, quién está "
               "disponible y quién pide ayuda, además de generar reportes gráficos "
               "(tiempo por día y por agente, horas pico y contactos más atendidos) y "
               "permitir la exportación de datos. La interfaz de la extensión se aísla "
               "mediante Shadow DOM para no interferir con WhatsApp Web, con el objetivo "
               "de reducir tiempos de respuesta y duplicidad de atención sin costos de "
               "licenciamiento."))
add_para(doc, "Componentes que integran el sistema:", bold=True)
add_numbered(doc, [
    ("Extensión de Chrome (cliente): popup de configuración y selección de agente, "
     "service worker con conexión WebSocket, content script que observa el DOM de "
     "WhatsApp Web, detector de contacto activo y panel flotante en Shadow DOM con "
     "lista de agentes, timers de chat y botones de pausa y ayuda."),
    ("Servidor de comunicaciones: punto de entrada HTTP + WebSocket en Node.js, "
     "administrador de agentes y presencia (RoomManager) con heartbeat, detección "
     "de desconexiones y reconexión de perfiles, y capa de persistencia SQLite."),
    ("Portal web de monitoreo (dashboard): interfaz servida por el propio servidor "
     "con las pestañas En vivo (estado actual de los agentes con actualización "
     "automática), Historial (sesiones por agente y contacto), Reportes (gráficas de "
     "tiempo por día, horas pico y top contactos con Chart.js) y Gestión (información "
     "del servidor y exportación de datos en JSON)."),
])

# ==================== d) OBJETIVOS ESPECÍFICOS =======================
add_heading(doc, "d) Objetivos específicos", size=14)
add_numbered(doc, [
    "Diseñar y construir una extensión de Google Chrome (Manifest V3) que se integre de manera no intrusiva con WhatsApp Web mediante un panel flotante con Shadow DOM.",
    "Implementar un servidor Node.js con comunicación WebSocket que administre el registro de agentes, sus estados de presencia y la limpieza de conexiones inactivas.",
    "Desarrollar un portal web de monitoreo (dashboard) que muestre en tiempo real el estado de los agentes y ofrezca módulos de Historial, Reportes y Gestión.",
    "Implementar la persistencia de la información en una base de datos SQLite con sesiones de atención (duración) y bitácora de eventos (chat_start, chat_end, paused, help_request, etc.).",
    "Generar métricas de desempeño a partir de los datos almacenados: tiempo atendido por día y por agente, horas pico de atención y contactos más atendidos.",
    "Permitir la gestión de agentes (agregar y eliminar) y el control manual del estado (pausar y reanudar) tanto desde el popup como desde el panel flotante.",
    "Garantizar la persistencia del perfil del agente y la reconexión automática con reintentos exponenciales ante cortes de red.",
    "Validar el sistema mediante pruebas funcionales con un equipo real de atención al cliente, incluyendo la solicitud de ayuda entre agentes.",
])

# ============ e) ANTECEDENTES Y JUSTIFICACIÓN ========================
add_heading(doc, "e) Antecedentes y justificación de la problemática a atender", size=14)
add_para(doc, "Antecedentes:", bold=True)
add_para(doc, ("En los equipos de servicio al cliente que atienden por WhatsApp no "
               "existe una herramienta nativa que indique qué agente está atendiendo "
               "qué conversación, quién está disponible o quién se encuentra pausado. "
               "Las soluciones existentes, como la WhatsApp Business API, suelen "
               "requerir costos de licenciamiento o desarrollo adicional y no ofrecen "
               "una vista de presencia de los agentes en tiempo real sobre la misma "
               "interfaz de WhatsApp Web."))
add_para(doc, "Justificación:", bold=True)
add_para(doc, ("La falta de visibilidad sobre el estado de cada agente provoca "
               "respuestas duplicadas, largos tiempos de espera del cliente y una "
               "distribución desigual de la carga de trabajo. El presente proyecto "
               "propone una solución de bajo costo que, al operar directamente sobre "
               "WhatsApp Web, no requiere de licencias adicionales ni de migrar la "
               "plataforma de atención. Esto permite que el equipo conozca en todo "
               "momento quién está atendiendo cada chat, quién está disponible y quién "
               "requiere ayuda, mejorando la coordinación, la calidad del servicio y la "
               "satisfacción del cliente. Además, al integrar un portal web de "
               "monitoreo, la administración dispone de métricas de desempeño (tiempos "
               "de atención, horas pico y contactos más atendidos) que respaldan la "
               "toma de decisiones."))
add_para(doc, "Beneficios esperados:", bold=True)
add_bullets(doc, [
    "Reducción de respuestas duplicadas y de los tiempos de espera del cliente.",
    "Distribución más equilibrada de la carga de trabajo entre los agentes.",
    "Coordinación eficiente mediante solicitudes de ayuda y visibilidad de estado en tiempo real.",
    "Obtención de métricas y reportes de desempeño para la administración.",
    "Bajo costo de implementación, sin licencias adicionales, cuando se opera sobre WhatsApp Web.",
])

# ============ f) DESCRIPCIÓN DETALLADA DE ACTIVIDADES ================
add_heading(doc, "f) Descripción detallada de las actividades (Metodología)", size=14)
add_para(doc, ("El desarrollo del proyecto se realizará bajo un enfoque iterativo e "
               "incremental, con las siguientes etapas:"))
add_numbered(doc, [
    ("Análisis de requerimientos: estudio del flujo de atención en WhatsApp Web, "
     "identificación de necesidades del equipo y definición de los estados de presencia."),
     ("Diseño de la arquitectura: definición de los componentes (extensión, content script, "
      "service worker, servidor WebSocket y portal de monitoreo) y del protocolo de "
      "comunicación entre ellos."),
    ("Configuración del entorno: instalación de Node.js 18+, inicialización del proyecto "
     "con Vite y TypeScript, y control de versiones con Git."),
    ("Desarrollo del servidor: implementación del servicio WebSocket con registro de "
     "agentes, administración de presencia, heartbeat y difusión de actualizaciones."),
    ("Desarrollo de la extensión: creación del popup de configuración, del content script "
     "con el panel en Shadow DOM, del detector de contacto activo y de la observación "
     "automática del chat que abre cada agente."),
    ("Implementación de la persistencia y reconexión: almacenamiento del perfil del agente "
     "y reconexión exponencial ante fallas de red."),
    ("Desarrollo de la base de datos: implementación del registro de sesiones y eventos "
     "en SQLite, con auto-guardado en disco."),
    ("Desarrollo del portal web de monitoreo: creación del dashboard con las pestañas En "
     "vivo, Historial, Reportes (gráficas con Chart.js) y Gestión, así como la ruta de "
     "exportación de datos en JSON."),
    ("Integración y pruebas del protocolo: pruebas de conexión, cambio de estados, "
     "eliminación de agentes, solicitudes de ayuda y desconexiones."),
    ("Pruebas funcionales: validación con un equipo real de atención al cliente para "
     "verificar usabilidad y comportamiento en condiciones de uso normal."),
    ("Documentación: elaboración de manuales, del reporte preliminar y del reporte final del proyecto."),
])

# ==================== g) CRONOGRAMA PRELIMINAR ========================
add_heading(doc, "g) Cronograma preliminar de actividades", size=14)
add_para(doc, (f"Duración mínima del proyecto: 4 meses, del {formato_fecha(FECHA_INICIO_PROYECTO)} "
               f"al {formato_fecha(FECHA_FIN_PROYECTO)}."), italic=True)

gantt_tmp = os.path.join(tempfile.gettempdir(), "wts_gantt.png")
build_gantt(ACTIVIDADES_CRONOGRAMA, COLORES_FASE, gantt_tmp)
pic = doc.add_paragraph()
pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic.add_run().add_picture(gantt_tmp, width=Inches(6.4))
add_para(doc, "", space_after=2)

headers = ["No.", "Actividad", "Fecha inicial", "Fecha final"]
rows = [(a[0].split(".")[0], a[0].split(". ", 1)[1], formato_fecha(a[1]), formato_fecha(a[2]))
        for a in ACTIVIDADES_CRONOGRAMA]
make_table(doc, headers, rows, widths=[1.0, 11.0, 2.3, 2.3])
add_para(doc, ("El proyecto dura 18 semanas (4 meses): cada actividad ocupa una semana y "
               "las actividades son consecutivas, iniciando la siguiente al terminar la "
               "anterior. Se organiza en fases: Análisis (1-2), Diseño y configuración "
               "(3-4), Desarrollo (5-10), Integración y pruebas (11-14) y Documentación y "
               "cierre (15-18). El cronograma inicia el viernes 7 de agosto de 2026 y "
               "concluye el 10 de diciembre de 2026."),
         size=9, italic=True, space_after=12)

# ================ h) RECURSOS MATERIALES REQUERIDOS ==================
add_heading(doc, "h) Recursos materiales requeridos", size=14)
make_table(doc, ["Recurso", "Tipo", "Cantidad", "Costo estimado"],
           RECURSOS, widths=[8.5, 4.0, 2.0, 3.0])
add_para(doc, "", space_after=2)
p = doc.add_paragraph()
set_font(p.add_run("Costo total estimado: "), size=11, bold=True)
set_font(p.add_run(recursos_totales), size=11)

# ================== i) INFORMACIÓN DE LA EMPRESA =====================
add_heading(doc, "i) Información de la empresa, organismo o dependencia", size=14)
add_bullets(doc, [
    f"Nombre: {EMPRESA['nombre']}",
    f"Dirección: {EMPRESA['direccion']}",
    f"Teléfono: {EMPRESA['telefono']}",
    f"Correo electrónico: {EMPRESA['correo']}",
    f"Giro: {EMPRESA['giro']}",
    f"Ciudad: {EMPRESA['ciudad']}",
])

# ================ j) LUGAR DONDE SE REALIZARÁ ========================
add_heading(doc, "j) Lugar donde se realizará el proyecto", size=14)
add_para(doc, "Ubicación geográfica o departamento: " + LUGAR_PROYECTO)

# ================== k) DATOS DEL ASESOR EXTERNO ======================
add_heading(doc, "k) Datos del asesor externo", size=14)
add_bullets(doc, [
    f"Departamento adscrito: {ASESOR['departamento']}",
    f"Puesto: {ASESOR['puesto']}",
    f"Correo electrónico: {ASESOR['correo']}",
    f"Teléfono: {ASESOR['telefono']}",
])

add_para(doc, "", space_after=8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("ITH-AC-PO-007-07  Rev."), size=10, bold=True, color=(0x66, 0x66, 0x66))

doc.save(DOC_OUTPUT)
print("OK ->", DOC_OUTPUT)
